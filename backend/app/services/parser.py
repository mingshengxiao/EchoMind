import io
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

from app.config import Settings

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}


class ParsedResume:
    def __init__(self, text: str, filename: str, file_size: int) -> None:
        self.text = text
        self.filename = filename
        self.file_size = file_size
        self.word_count = len(text.split())
        self.preview = text[:800]


def _extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(part.strip() for part in parts if part.strip())


def _parse_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text.strip())
    return "\n".join(paragraphs + table_cells)


def parse_resume_bytes(filename: str, data: bytes, settings: Settings) -> ParsedResume:
    extension = _extension(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file type. Allowed: PDF, DOCX, MD, TXT.",
        )
    if len(data) > settings.max_upload_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Max: {settings.max_upload_mb}MB.",
        )
    if not data:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")

    try:
        if extension == ".pdf":
            text = _parse_pdf(data)
        elif extension == ".docx":
            text = _parse_docx(data)
        else:
            text = _decode_text(data)
    except (BadZipFile, ValueError, OSError) as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unable to parse this resume file") from exc

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(normalized.split()) < 10:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Resume does not contain enough readable text")

    return ParsedResume(text=normalized, filename=filename, file_size=len(data))


async def parse_upload(file: UploadFile, settings: Settings) -> ParsedResume:
    data = await file.read()
    return parse_resume_bytes(file.filename or "resume", data, settings)
